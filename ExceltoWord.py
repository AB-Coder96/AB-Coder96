import openpyxl
from openpyxl import Workbook
from openpyxl.styles import Alignment
from docx import Document

# Create Excel workbook and add data to "skills" sheet
wb = Workbook()
ws = wb.active
ws.title = "skills"

# Sample data for the first two columns
data_col1_col2 = [
    ["Skill 1", "Description 1"],
    ["Skill 2", "Description 2"],
    # Add more rows as needed
]

# Populate the first two columns
for row in data_col1_col2:
    ws.append(row)

# Create additional sheets for "projects" and populate tables
for i in range(3, 7, 2):
    # Sample data for columns i, i+1
    data_col_i_i1 = [
        [f"Project {i}", f"Description {i}"],
        [f"Project {i+1}", f"Description {i+1}"],
        # Add more rows as needed
    ]

    # Create a new sheet for each pair of columns
    ws_proj = wb.create_sheet(title=f"projects_{i}_{i+1}")

    # Populate the table in the new sheet
    for row in data_col_i_i1:
        ws_proj.append(row)

# Save the Excel workbook
wb.save("projects.xlsx")

# Create Word document and add content
doc = Document()

# Add header "Skills" to Word document
doc.add_heading("Skills", level=1)

# Add the first two columns from the "skills" sheet to the Word document
for row in data_col1_col2:
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    for j, cell_data in enumerate(row):
        cell = table.cell(0, j)
        cell.text = cell_data

# Add headers and tables from each "projects" sheet to the Word document
for i in range(3, 7, 2):
    sheet_name = f"projects_{i}_{i+1}"
    doc.add_heading(f"Full List of Projects - {sheet_name}", level=1)

    # Retrieve data from the corresponding sheet
    ws_proj = wb[sheet_name]
    data_proj = [list(row) for row in ws_proj.iter_rows(values_only=True)]

    # Add table to Word document
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    for row in data_proj:
        table.add_row()
        for j, cell_data in enumerate(row):
            cell = table.cell(table.rows.index(table.last_row), j)
            cell.text = cell_data

# Save the Word document
doc.save("portfolio.docx")

print("Portfolio created successfully.")
