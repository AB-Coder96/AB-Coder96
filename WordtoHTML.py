from docx import Document
from datetime import datetime

def extract_paragraphs_and_tables(docx_file):
    doc = Document(docx_file)
    paragraphs = [p.text for p in doc.paragraphs]
    tables_html = []

    for table in doc.tables:
        table_html = "<table border='1'>"
        
        # Add the header row with color shading
        table_html += "<tr style='background-color: #808080;'>"
        for cell in table.rows[0].cells:
            table_html += f"<th>{cell.text}</th>"
        table_html += "</tr>"
        
        # Add the remaining rows
        for row in table.rows[1:]:
            table_html += "<tr>"
            for cell in row.cells:
                table_html += f"<td>{cell.text}</td>"
            table_html += "</tr>"
        
        table_html += "</table>"
        tables_html.append(table_html)

    return paragraphs, tables_html
def generate_html_css_from_docx(docx_file, output_html_file):
    paragraphs, tables = extract_paragraphs_and_tables(docx_file)

    print("Paragraphs:")
    for i, paragraph in enumerate(paragraphs):
        print(f"Paragraph {i + 1}: {paragraph}")

    print("\nTables:")
    for i, table in enumerate(tables):
        print(f"Table {i + 1}:\n{table}")

    html_content = "<!DOCTYPE html>\n<html>\n<head>\n<title>Portfolio</title>\n"
    html_content += "<link rel='stylesheet' href='https://cdnjs.cloudflare.com/ajax/libs/font-awesome/5.15.2/css/all.min.css' integrity='sha384-rGp68ppUyMMd6JAiyr5K1SfSvYY4GYvRnfT2XrSP7z2S+YYZxM3R8yf+BxyiXagD' "
    html_content += "crossorigin='anonymous'>\n"
    html_content += "<style>\nbody {\nfont-family: Arial, sans-serif;\nmargin: 20px;\n}\n"
    html_content += "table {\nborder-collapse: collapse;\nmargin: 0 auto;\n}\ntable, th, td {\nborder: 1px solid black;\npadding: 5px;\n}\n"
    html_content += "h1, p {\ntext-align: center;\n}\n"
    html_content += "footer {\nposition: fixed;\nbottom: 0;\nwidth: 100%;\nbackground-color: #333;\ntext-align: center;\npadding: 10px;\n"
    html_content += "color: white;\n}\n"
    html_content += "footer a {\ncolor: white;\n}\n"
    html_content += ".table-container {\n display: flex;\n }\n"
    html_content += ".table-container table {\n flex: 1;\n margin-right: 20px;\n}\n" 
    html_content += "</style>\n</head>\n<body>\n"


    # Add the portion showing experience and projects completed
    html_content += "<style>\n"
    html_content += "@import url('https://fonts.googleapis.com/css2?family=Roboto:wght@700&display=swap');\n"
    html_content += "</style>\n"

    html_content += "<div style='text-align: center;'>\n"
    html_content += "    <div style='display: inline-block; background-color: #f0f8ff; padding: 10px; border-radius: 5px; font-family: Roboto, sans-serif;'>\n"
    html_content += "        <div style='font-size: 20px; color: #00008b;'>Years of Experience:</div>\n"
    html_content += "        <div class='counter' id='years-counter' style='font-size: 24px; color: #00008b;'>0</div>\n"
    html_content += "    </div>\n"
    html_content += "    <div style='display: inline-block; background-color: #f0f8ff; padding: 10px; border-radius: 5px; font-family: Roboto, sans-serif;'>\n"
    html_content += "        <div style='font-size: 20px; color: #00008b;'>Projects Completed:</div>\n"
    html_content += "        <div class='counter' id='projects-counter' style='font-size: 24px; color: #00008b;'>0</div>\n"
    html_content += "    </div>\n"
    html_content += "</div>\n"

    html_content += "<script>\n"
    html_content += "    const yearsCounter = document.getElementById('years-counter');\n"
    html_content += "    const projectsCounter = document.getElementById('projects-counter');\n"
    html_content += "    let years = 0;\n"
    html_content += "    let projects = 0;\n"
    html_content += "    const yearsIncrement = 4 / 200; // Decreased increment for slower animation\n"
    html_content += "    const projectsIncrement = 23 / 200; // Decreased increment for slower animation\n"
    html_content += "    function animateCounters() {\n"
    html_content += "        if (years < 4) {\n"
    html_content += "            years += yearsIncrement;\n"
    html_content += "            yearsCounter.innerText = Math.round(years);\n"
    html_content += "        }\n"
    html_content += "        if (projects < 23) {\n"
    html_content += "            projects += projectsIncrement;\n"
    html_content += "            projectsCounter.innerText = Math.round(projects);\n"
    html_content += "        }\n"
    html_content += "        requestAnimationFrame(animateCounters);\n"
    html_content += "    }\n"
    html_content += "    animateCounters();\n"
    html_content += "</script>\n"




    # Add the first paragraph as an <h1> header
    if paragraphs:
        html_content += f"<h1>{paragraphs[0]}</h1>\n"
    html_content += f"<p>{paragraphs[1]}</p>\n"
    html_content += "<div class='table-container'>\n"
    html_content += tables[0] + "\n"
    html_content += tables[1] + "\n"
    html_content += "</div>\n"
    # Check if the current paragraph is at index 4
    html_content += f"<p>{paragraphs[4]}</p>\n"
    # Display the table
    html_content += "<div class='table-container'>\n"
    html_content += tables[2] + "\n"
    html_content += tables[3] + "\n"
    html_content += "</div>\n"
    html_content += f"<p>{paragraphs[6]}</p>\n"
    # Display the table
    html_content += "<div class='table-container'>\n"
    html_content += tables[4] + "\n"
    html_content += tables[5] + "\n"
    html_content += "</div>\n"
    html_content += f"<h1>{paragraphs[9]}</h1>\n"
    html_content += tables[6] + "\n"
    # Add code to embed PDF using PDF.js
    html_content += "<canvas id='pdf-render'></canvas>\n"
    html_content += "<script src='https://mozilla.github.io/pdf.js/build/pdf.js'></script>\n"
    html_content += "<script>\n"
    html_content += "const pdfUrl = 'Imbalance_paper.pdf';\n"  # Update this line with the correct PDF file name
    html_content += "pdfjsLib.getDocument(pdfUrl).then(pdf => {\n"
    html_content += "pdf.getPage(1).then(page => {\n"
    html_content += "const canvas = document.getElementById('pdf-render');\n"
    html_content += "const context = canvas.getContext('2d');\n"
    html_content += "const viewport = page.getViewport({ scale: 1.5 });\n"
    html_content += "canvas.width = viewport.width;\n"
    html_content += "canvas.height = viewport.height;\n"
    html_content += "page.render({ canvasContext: context, viewport: viewport });\n"
    html_content += "});\n"
    html_content += "});\n"
    html_content += "</script>\n"


    # Add HTML footer with date modified, links, and email
    date_modified = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    github_link = "https://github.com/AB-Coder96"
    linkedin_link = "https://www.linkedin.com/in/araz-karimi-0b2600290/"
    email_address = "arazbagherzadeh@gmail.com"
    
    html_content += f"<footer>\n"
    html_content += f"Made by Araz Karimi<br>\n"  # Line added before "Date Modified"
    html_content += f"<a href='{github_link}' target='_blank'><i class='fab fa-github'></i> GitHub</a> | "
    html_content += f"<a href='{linkedin_link}' target='_blank'><i class='fab fa-linkedin'></i> LinkedIn</a> | "
    html_content += f"<a href='mailto:{email_address}'><i class='far fa-envelope'></i> {email_address}</a><br>\n"
    html_content += f"Date Modified: {date_modified}\n"
    html_content += "</footer>\n"

    html_content += "</body>\n</html>"

    with open(output_html_file, "w", encoding="utf-8") as html_file:
        html_file.write(html_content)

# Replace 'input_word.docx' with the path to your Word file
input_word_file = 'Araz B Karimi Portfolio.docx'
output_html_file = 'docs/index.html'

generate_html_css_from_docx(input_word_file, output_html_file)